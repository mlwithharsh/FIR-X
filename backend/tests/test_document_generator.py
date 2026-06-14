from io import BytesIO
from pathlib import Path

from docx import Document

from app.schemas.report import DARForm
from app.services.document_generator import DocumentGenerator


TEMPLATE_PATH = Path(__file__).resolve().parents[1] / "app" / "template_dar.docx"


def _build_form() -> DARForm:
    values = {
        field_name: "TEST-VALUE"
        for field_name, field in DARForm.model_fields.items()
        if field.is_required()
    }
    values.update(
        {
            "fir_no": "FIR-SENTINEL",
            "vehicle_reg_no": "REG-SENTINEL",
            "insurance_policy_no": "POLICY-SENTINEL",
            "insurance_period": "PERIOD-SENTINEL",
            "legal_representatives": [
                {
                    "name": "LOR-NAME-SENTINEL",
                    "relation": "LOR-RELATION-SENTINEL",
                    "age": "LOR-AGE-SENTINEL",
                    "gender": "LOR-GENDER-SENTINEL",
                    "marital_status": "LOR-MARITAL-SENTINEL",
                    "address": "LOR-ADDRESS-SENTINEL",
                    "contact": "LOR-CONTACT-SENTINEL",
                }
            ],
            "minor_children": [
                {
                    "name": "CHILD-NAME-SENTINEL",
                    "school_class": "CHILD-SCHOOL-SENTINEL",
                    "annual_fee": "CHILD-FEE-SENTINEL",
                    "approximate_expenses": "CHILD-EXPENSE-SENTINEL",
                }
            ],
        }
    )
    return DARForm(**values)


def _generate_document() -> Document:
    content = DocumentGenerator(TEMPLATE_PATH).generate_dar(_build_form())
    return Document(BytesIO(content))


def _table_text(document: Document, table_index: int) -> str:
    return "\n".join(
        cell.text
        for row in document.tables[table_index].rows
        for cell in row.cells
    )


def test_reported_policy_period_and_registration_fields_are_parameterized():
    document = _generate_document()

    for table_index in (22, 25, 28, 30, 38):
        table_text = _table_text(document, table_index)
        assert "POLICY-SENTINEL" in table_text
        assert "PERIOD-SENTINEL" in table_text
        assert "REG-SENTINEL" in table_text


def test_form_vi_a_legal_representative_contact_row_uses_frontend_values():
    document = _generate_document()
    table_text = _table_text(document, 31)

    assert "LOR-NAME-SENTINEL" in table_text
    assert "LOR-CONTACT-SENTINEL" in table_text
    assert "LOR-ADDRESS-SENTINEL" in table_text
    assert "LOR-GENDER-SENTINEL" in table_text
    assert "LOR-MARITAL-SENTINEL" in table_text


def test_form_vi_a_minor_child_and_under_section_fields_render():
    document = _generate_document()

    assert "CHILD-NAME-SENTINEL" in _table_text(document, 31)
    assert "CHILD-SCHOOL-SENTINEL" in _table_text(document, 31)
    assert "281/106(1) BNS, 2023" in _table_text(document, 0)


def test_existing_uppercase_template_placeholders_still_render():
    document = _generate_document()
    document_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )

    assert "FIR-SENTINEL" in document_text
    assert "{{" not in document_text
    assert "{%" not in document_text
