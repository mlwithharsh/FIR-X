from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.section import Section
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from app.schemas.report import LegalRepresentativeSchema, ReportRequest
from app.services.conditional_logic import build_conditional_sections


class DocumentGenerator:
    _TEMPLATE_NAME = "dar_reference.docx"

    def __init__(self, template_dir: Path, generated_dir: Path) -> None:
        self.template_dir = template_dir
        self.generated_dir = generated_dir
        self.generated_dir.mkdir(parents=True, exist_ok=True)

    def generate_dar(self, data: ReportRequest, legal_description: str, filename: str) -> Path:
        template_path = self.template_dir / self._TEMPLATE_NAME
        if not template_path.exists():
            raise FileNotFoundError(f"DAR template not found: {template_path}")

        document = Document(template_path)
        replacements = self._build_replacements(data, legal_description)
        self._apply_replacements(document, replacements)

        output_path = self.generated_dir / filename
        document.save(output_path)
        return output_path

    def _apply_replacements(self, document: DocumentObject, replacements: dict[str, str]) -> None:
        paragraphs = list(self._iter_document_paragraphs(document))
        for old, new in replacements.items():
            if not old:
                continue
            for paragraph in paragraphs:
                self._replace_in_paragraph(paragraph, old, new)

    def _iter_document_paragraphs(self, document: DocumentObject) -> Iterable[Paragraph]:
        yield from self._iter_paragraphs_from_parent(document)
        for section in document.sections:
            yield from self._iter_section_paragraphs(section)

    def _iter_section_paragraphs(self, section: Section) -> Iterable[Paragraph]:
        for part in (section.header, section.first_page_header, section.even_page_header, section.footer, section.first_page_footer, section.even_page_footer):
            yield from self._iter_paragraphs_from_parent(part)

    def _iter_paragraphs_from_parent(self, parent: DocumentObject | _Cell | Table) -> Iterable[Paragraph]:
        if hasattr(parent, "paragraphs"):
            for paragraph in parent.paragraphs:
                yield paragraph
        if hasattr(parent, "tables"):
            for table in parent.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from self._iter_paragraphs_from_parent(cell)

    def _replace_in_paragraph(self, paragraph: Paragraph, old: str, new: str) -> None:
        if not paragraph.runs:
            if old in paragraph.text:
                paragraph.text = paragraph.text.replace(old, new)
            return

        while old in "".join(run.text for run in paragraph.runs):
            full_text = "".join(run.text for run in paragraph.runs)
            start = full_text.find(old)
            if start < 0:
                return
            end = start + len(old)

            run_positions: list[tuple[int, int, int]] = []
            cursor = 0
            for run_index, run in enumerate(paragraph.runs):
                next_cursor = cursor + len(run.text)
                run_positions.append((run_index, cursor, next_cursor))
                cursor = next_cursor

            first_index = 0
            last_index = 0
            first_offset = 0
            last_offset = 0
            for run_index, run_start, run_end in run_positions:
                if run_start <= start < run_end:
                    first_index = run_index
                    first_offset = start - run_start
                if run_start < end <= run_end:
                    last_index = run_index
                    last_offset = end - run_start
                    break

            first_run = paragraph.runs[first_index]
            last_run = paragraph.runs[last_index]
            prefix = first_run.text[:first_offset]
            suffix = last_run.text[last_offset:]
            first_run.text = prefix + new + suffix
            for run_index in range(first_index + 1, last_index + 1):
                paragraph.runs[run_index].text = ""

    def _build_replacements(self, data: ReportRequest, legal_description: str) -> dict[str, str]:
        case = data.case_details
        conditional = build_conditional_sections(data)
        plaintiff_name = case.plaintiff_name or (data.legal_representatives[0].name if data.legal_representatives else data.victim.name)
        plaintiff_age = case.plaintiff_age or (data.legal_representatives[0].age if data.legal_representatives else data.victim.age or "")
        plaintiff_mobile = case.plaintiff_mobile or ""
        defendant_name = case.defendant_name or data.driver.name
        defendant_age = case.defendant_age or data.driver.age or ""
        defendant_mobile = case.defendant_mobile or data.driver.phone or ""
        fir_date = case.fir_date.strftime("%d/%m/%Y")
        fir_date_short = case.fir_date.strftime("%d/%m/%y")
        accident_date = data.accident.date.strftime("%d/%m/%Y")
        accident_time = f"At about {data.accident.time.strftime('%H:%M')} hrs"
        long_sections = case.sections if "2023" in case.sections else f"{case.sections}, 2023"
        police_station_dot = self._with_dot(case.police_station)
        district = case.district or ""
        district_sentence = f"Distt {district}." if district else "Distt __________."
        io_name = case.investigating_officer
        io_phone = case.investigating_officer_phone or ""
        owner_address = data.vehicle.owner_address or ""
        driver_address = data.driver.address or ""
        victim_address = data.victim.address or ""
        victim_age = data.victim.age or ""
        victim_line = f"{data.victim.name}{self._optional_prefix(' R/o ', victim_address)}{self._optional_prefix('. Age-', victim_age)}"
        owner_line = f"Offending vehicle registered owner- {data.vehicle.owner_name}{self._optional_prefix(' R/o ', owner_address)}{self._optional_prefix('. Mb No. ', data.vehicle.owner_phone)}."
        driver_line = f"Accused driver- {data.driver.name}{self._optional_prefix(' R/o ', driver_address)}{self._optional_prefix('. Age-', data.driver.age)}{self._optional_prefix(' Mb No. ', data.driver.phone)}"
        hospital_name = data.hospital.name or ""
        hospital_address = data.hospital.address or ""
        doctor_name = data.hospital.doctor_name or ""
        replacements = {
            "212/2026": case.fir_number,
            "11/03/2026": fir_date,
            "11/03/26": fir_date_short,
            "281/106(1) BNS, 2023": long_sections,
            "281/106(1) BNS": case.sections,
            "SP Badli Delhi.": police_station_dot,
            "SP Badli_Distt Outer North Delhi.": f"{case.police_station}_{district_sentence}".strip("_"),
            "SP Badli, Delhi": case.police_station,
            "Distt Outer North Delhi.": district_sentence,
            "ASI Satyaveer No.6268-D": io_name,
            "MACT-OND, Delhi.": district or "MACT-OND, Delhi.",
            "MACT- OND Delhi.": district or "MACT- OND Delhi.",
            "Case FIR No.212/2026 dt.11/03/2026 u/s 281/106(1) BNS PS SP Badli, Delhi.": f"Case FIR No.{case.fir_number} dt.{fir_date} u/s {case.sections} PS {case.police_station}.",
            "Deceased of the case- Mrs. Prem Wati W/o Bakshi Singh R/o J-491, Bhagwan Pura Samaypur Libaspur Delhi. Age-84 Yrs": f"{'Deceased' if data.victim.status == 'deceased' else 'Injured'} of the case- {victim_line}",
            "Offending vehicle registered owner- Ramesh Chand S/o Bidhi Chand R/o HNo.233, Ward No.3 Khera Sita Ram Kalka Panchkula Haryana. Mb No. 9816043050.": owner_line,
            "Accused driver- Babu Singh S/o Kundan Singh R/o Village Dabadi Ki Ser Chanyana Bakyori (257), Sirmor HP-173024. Age-50 Yrs Mb No. 9805392670": driver_line,
            "Offending vehicle insured- Chola MS General Insurance Co Ltd Delhi.": f"Offending vehicle insured- {data.insurance.company_name or 'Not supplied'}.",
            "At about 13:00 hrs": accident_time,
            "Infront of Tirpal Ghar Shop No CW-544, Cut of Sanjay Gandhi Transport Nagar Delhi.": data.accident.location,
            "Driver/Owner Victim Witness √Hospital Good Samaritan Police Others (Specify)": self._source_label(data),
            "Ct Kuldeep No.2219/NW": data.informant.name or "",
            "BJRM Hospital Jahangirpuri Delhi.": data.informant.address or hospital_address,
            "Injury √ Fatal Damage / loss of the property Any other loss/injury": self._nature_label(data),
            "One vehicle": self._number_of_vehicles_label(data.accident.number_of_vehicles),
            "√Yes No": self._yes_no_label(True),
            "Yes √ No": self._yes_no_label(data.accident.cctv_available),
            "One": self._count_label(data.accident.fatalities_count),
            "N/A": self._injured_count_label(data),
            "BJRM Hospital": hospital_name,
            "Jahangirpuri Delhi.": hospital_address,
            "Dr. Manish Kumar MO": doctor_name,
            "HR64A- 6664": data.vehicle.registration_number,
            "HR64A-6664": data.vehicle.registration_number,
            "Babu Singh S/o Kundan Singh": data.driver.name,
            "Village Dabadi Ki Ser Chanyana Bakyori (257), Sirmor HP-173024. Age-50 Yrs": self._compose_address_age(driver_address, data.driver.age),
            "9805392670": data.driver.phone or defendant_mobile,
            "Ramesh Chand S/o Bidhi Chand": data.vehicle.owner_name,
            "HNo.233, Ward No.3 Khera Sita Ram Kalka Panchkula Haryana.": owner_address,
            "9816043050.": self._with_dot(data.vehicle.owner_phone),
            "3379/04146458/000/01": data.insurance.policy_number or "",
            "30/11/25 to 29/11/26": data.insurance.policy_period or "",
            "Chola MS General Insurance Co Ltd Delhi.": self._with_dot(data.insurance.company_name),
            "Delhi.": self._with_dot(data.insurance.company_address),
            "Mrs. Prem Wati W/o Bakshi Singh": data.victim.name,
            "J-491, Bhagwan Pura Samaypur Libaspur Delhi. Age-84 Yrs": self._compose_address_age(victim_address, victim_age),
            "84 Yrs": victim_age,
            "Desi Vadi": data.victim.occupation or "",
            "Male √ Female Other": self._gender_label(data.driver.gender),
            "12,000/- PM": data.driver.monthly_income or "",
            "√Permanent Learner’s Juvenile Without License Others(Specify)": self._license_type_label(data.driver.license_type),
            "HP16A 20230000160": data.driver.license_number or "",
            "10/05/23 to 09/05/33": data.driver.license_validity or "",
            "Pachhad (HP16A)": data.driver.licensing_authority or "",
            "Plaintiff Name :________ Mukesh_______________________": f"Plaintiff Name :________ {plaintiff_name}_______________________",
            "Plaintiff Age :____________40 Years __________Male/Female": f"Plaintiff Age :____________{plaintiff_age} __________{self._gender_word(case.plaintiff_gender)}",
            "Plaintiff Mobile No :_________ 9136804031_________________": f"Plaintiff Mobile No :_________ {plaintiff_mobile}_________________",
            "Defendant Name :__________Babu Singh _____________": f"Defendant Name :__________{defendant_name} _____________",
            "Defendant Age :____________50 Yrs___________Male/Female": f"Defendant Age :____________{defendant_age}___________{self._gender_word(case.defendant_gender or data.driver.gender)}",
            "Defendant Mobile No :________ 9805392670_________________": f"Defendant Mobile No :________ {defendant_mobile}_________________",
            "Case FIR No & date :________212/2026_ dt.11/03/26___________": f"Case FIR No & date :________{case.fir_number}_ dt.{fir_date_short}___________",
            "Police Station :________SP Badli, Delhi_____________________": f"Police Station :________{case.police_station}_____________________",
            "U/S :_______281/106(1) BNS______________": f"U/S :_______{case.sections}______________",
            "Name of I.O. :_______ASI Satyaveer No.6268-D_________": f"Name of I.O. :_______{io_name}_________",
            "Mobile No. I.O. :_________ ______________________": f"Mobile No. I.O. :_________ {io_phone}______________________",
            "IN Case FIR No._212/2026_Date-11/03/2026 U/s-281/106(1) BNS PS SP Badli_Distt Outer North Delhi.": f"IN Case FIR No._{case.fir_number}_Date-{fir_date} U/s-{case.sections} PS {case.police_station}_{district_sentence}",
            "On 11/03/2026 at approximately 13:00 hours, at Sanjay Gandhi Transport Nagar, Delhi, the offending vehicle bearing registration number HR64A-6664, stated to be a Truck, allegedly driven by Babu Singh and owned by Ramesh Chand, was involved in a truck hit pedestrian. During the occurrence, Prem Wati, aged about 45 years, succumbed to the injuries. The matter requires consideration for compensation and compliance proceedings under the MACT framework. The vehicle is presently linked with insurance details not yet supplied.": legal_description,
        }

        for old, new in self._legal_representative_replacements(data.legal_representatives, victim_address):
            replacements[old] = new

        if conditional["has_insurance_details"] is False:
            replacements["Chola MS General Insurance Co Ltd Delhi."] = ""
            replacements["Delhi."] = ""

        return replacements

    def _legal_representative_replacements(self, representatives: list[LegalRepresentativeSchema], default_address: str) -> list[tuple[str, str]]:
        paragraph_templates = [
            ("Son of deceased- Mukesh S/o Lt. Bakshi Singh R/o J-491, Bhagwan Pura Samaypur Libaspur Delhi. Age-43 Yrs", 0),
            ("Son of deceased- Bhagat Singh S/o Lt. Bakshi Singh R/o J-491, Bhagwan Pura Samaypur Libaspur Delhi. Age-65 Yrs", 1),
            ("Daughter of deceased- Tilak Kaur D/o Lt. Bakshi Singh R/o J-491, Bhagwan Pura Samaypur Libaspur Delhi. Age-47 Yrs", 2),
        ]
        table_templates = [
            ("Mukesh S/o Lt. Bakshi Singh", "Son", "43 Yrs"),
            ("Bhagat Singh S/o Lt. Bakshi Singh", "Son", "65 Yrs"),
            ("Tilak Kaur D/o Lt. Bakshi Singh", "Daughter", "47 Yrs"),
        ]
        replacements: list[tuple[str, str]] = []
        for template, index in paragraph_templates:
            if index < len(representatives):
                lr = representatives[index]
                label = f"{lr.relation} of deceased".strip()
                address = lr.address or default_address
                text = f"{label}- {lr.name}{self._optional_prefix(' R/o ', address)}{self._optional_prefix('. Age-', lr.age)}"
                replacements.append((template, text))
            else:
                replacements.append((template, ""))

        for index, (name_sample, relation_sample, age_sample) in enumerate(table_templates):
            if index < len(representatives):
                lr = representatives[index]
                replacements.extend(
                    [
                        (name_sample, lr.name),
                        (relation_sample, lr.relation),
                        (age_sample, lr.age or ""),
                    ]
                )
            else:
                replacements.extend([(name_sample, ""), (relation_sample, ""), (age_sample, "")])
        return replacements

    def _compose_address_age(self, address: str, age: str | None) -> str:
        if not address and not age:
            return ""
        if address and age:
            return f"{address}. Age-{age}"
        return address or (age or "")

    def _source_label(self, data: ReportRequest) -> str:
        labels = {
            "driver_owner": "Driver/Owner",
            "victim": "Victim",
            "witness": "Witness",
            "hospital": "Hospital",
            "good_samaritan": "Good Samaritan",
            "police": "Police",
            "other": f"Others ({data.accident.other_source or 'Specify'})",
        }
        ordered = ["driver_owner", "victim", "witness", "hospital", "good_samaritan", "police", "other"]
        return " ".join(f"√{labels[item]}" if item == data.accident.source_of_information else labels[item] for item in ordered)

    def _nature_label(self, data: ReportRequest) -> str:
        return " ".join(
            [
                "√Injury" if data.accident.type == "injury" else "Injury",
                "√Fatal" if data.accident.type == "fatal" else "Fatal",
                "Damage / loss of the property",
                "Any other loss/injury",
            ]
        )

    def _yes_no_label(self, value: bool) -> str:
        return "√Yes No" if value else "Yes √No"

    def _number_of_vehicles_label(self, count: int) -> str:
        return "One vehicle" if count == 1 else f"{count} vehicles"

    def _count_label(self, count: int) -> str:
        names = {0: "N/A", 1: "One", 2: "Two", 3: "Three"}
        return names.get(count, str(count))

    def _injured_count_label(self, data: ReportRequest) -> str:
        return self._count_label(data.accident.injured_count) if data.accident.type == "injury" else "N/A"

    def _gender_label(self, gender: str | None) -> str:
        return " ".join(
            [
                "√Male" if gender == "male" else "Male",
                "√Female" if gender == "female" else "Female",
                "√Other" if gender == "other" else "Other",
            ]
        )

    def _gender_word(self, gender: str | None) -> str:
        if gender == "male":
            return "Male"
        if gender == "female":
            return "Female"
        if gender == "other":
            return "Other"
        return "Male/Female"

    def _license_type_label(self, license_type: str | None) -> str:
        mapping = {
            "permanent": "Permanent",
            "learner": "Learner’s",
            "juvenile": "Juvenile",
            "without_license": "Without License",
            "other": "Others(Specify)",
        }
        ordered = ["permanent", "learner", "juvenile", "without_license", "other"]
        return " ".join(f"√{mapping[item]}" if item == license_type else mapping[item] for item in ordered)

    def _optional_prefix(self, prefix: str, value: str | None) -> str:
        return f"{prefix}{value}" if value else ""

    def _with_dot(self, value: str | None) -> str:
        if not value:
            return ""
        return value if value.endswith(".") else f"{value}."
