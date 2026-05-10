from docx import Document

def search_text(docx_path, search_str):
    doc = Document(docx_path)
    found = False
    for para in doc.paragraphs:
        if search_str in para.text:
            print(f"Found '{search_str}' in paragraph: {para.text}")
            found = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if search_str in cell.text:
                    print(f"Found '{search_str}' in table cell: {cell.text}")
                    found = True
    return found

if __name__ == "__main__":
    template_path = r"d:\FIR-Auto\backend\app\template_dar.docx"
    print("Searching for default accident time '13:00'...")
    search_text(template_path, "13:00")
    print("\nSearching for default accident place...")
    search_text(template_path, "Sanjay Gandhi Transport Nagar")
