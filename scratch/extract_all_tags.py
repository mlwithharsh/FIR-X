import re
from docx import Document

def extract_all_potential_tags(docx_path):
    doc = Document(docx_path)
    tags = set()
    
    # Regex to find {{ ... }}
    pattern = re.compile(r'\{\{\s*(.*?)\s*\}\}')
    
    def find_in_text(text):
        if text:
            matches = pattern.findall(text)
            for m in matches:
                tags.add(m)

    for para in doc.paragraphs:
        find_in_text(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                find_in_text(cell.text)
                for p in cell.paragraphs:
                    find_in_text(p.text)
                    
    return tags

if __name__ == "__main__":
    template_path = r"d:\FIR-Auto\backend\app\template_dar.docx"
    tags = extract_all_potential_tags(template_path)
    print("All tags found in template:")
    for t in sorted(list(tags)):
        print(t)
