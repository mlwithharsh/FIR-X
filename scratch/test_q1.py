import requests
import json
import os
from docx import Document
from io import BytesIO

BASE_URL = "http://127.0.0.1:8000"

payload = {
  "fir_no": "TEST-9999",
  "date": "10/05/2026",
  "accident_date": "10/05/2026",
  "accident_time": "14:00",
  "accident_place": "QA Place Sentinel",
  "driver_name": "QADRIVER_SENTINEL",
  "driver_father": "QA_FATHER",
  "driver_address": "QA_ADDRESS",
  "driver_mobile": "9999999999",
  "driver_dl_no": "QA_DL_123",
  "driver_dl_validity": "2030",
  "driver_dl_authority": "QA_AUTH",
  "owner_name": "QAOWNER_SENTINEL",
  "owner_father": "QA_OWNER_FATHER",
  "owner_address": "QA_OWNER_ADDRESS",
  "owner_mobile": "8888888888",
  "vehicle_reg_no": "QA00AA0000",
  "vehicle_make": "QA_MAKE",
  "vehicle_model": "QA_MODEL",
  "vehicle_year": "2024",
  "vehicle_chassis": "QA_CHASSIS",
  "vehicle_engine": "QA_ENGINE",
  "insurance_policy_no": "QA_INS_123",
  "insurance_period": "2024-2025",
  "insurance_co_name": "QA_INS_CO",
  "insurance_co_address": "QA_INS_ADDR",
  "deceased_name": "QA_VICTIM",
  "deceased_husband": "QA_HUSBAND",
  "deceased_age": "30",
  "deceased_address": "QA_VICTIM_ADDR",
  "deceased_occupation": "QA_JOB",
  "deceased_income": "50000",
  "hospital_name": "QA_HOSPITAL",
  "hospital_address": "QA_HOSP_ADDR",
  "doctor_name": "QA_DOCTOR",
  "io_name": "QA_IO",
  "io_rank": "QA_RANK",
  "io_pis": "QA_PIS",
  "io_phone": "7777777777",
  "ps_name": "QA_PS",
  "legal_representatives": []
}

def test_t1():
    print("Running T1...")
    response = requests.post(f"{BASE_URL}/api/generate", json=payload)
    if response.status_code != 200:
        print(f"FAIL: T1 - Status code {response.status_code}")
        return

    doc = Document(BytesIO(response.content))
    all_text = []
    for para in doc.paragraphs:
        all_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text.append(cell.text)
    
    full_text = "\n".join(all_text)
    
    if "QADRIVER_SENTINEL" in full_text:
        print("PASS: T1 - Sentinel found")
    else:
        if "Babu Singh" in full_text:
            print("FAIL: T1 - Original sample name 'Babu Singh' still appears")
        else:
            print("FAIL: T1 - Sentinel NOT found, but 'Babu Singh' also not found. Check document structure.")

def test_t2():
    print("Running T2...")
    payload2 = payload.copy()
    payload2["fir_no"] = "TEST-8888"
    payload2["driver_name"] = "OTHER_DRIVER"
    
    res1 = requests.post(f"{BASE_URL}/api/generate", json=payload)
    res2 = requests.post(f"{BASE_URL}/api/generate", json=payload2)
    
    if res1.content == res2.content:
        print("FAIL: T2 - Both files are identical")
    else:
        print("PASS: T2 - Files differ")

def test_t3():
    print("Running T3...")
    payload3 = payload.copy()
    # The schema requires all fields to be min_length=1. 
    # Let's check if we can leave one "blank" (empty string).
    # Wait, the schema says Field(min_length=1). So empty string might fail validation.
    # The test says "Leave one non-required field blank".
    # I should check if there are non-required fields.
    # In report.py, legal_representatives is the only one with a default.
    # Others have min_length=1.
    
    # Try sending an empty string for vehicle_model if it allows it.
    # Actually, I'll check the schema again.
    pass

if __name__ == "__main__":
    test_t1()
    test_t2()
