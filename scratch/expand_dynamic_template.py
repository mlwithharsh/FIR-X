from pathlib import Path

from docx import Document


TEMPLATE = Path(__file__).resolve().parents[1] / "backend" / "app" / "template_dar.docx"


def set_cell(table, row: int, cell: int, value: str) -> None:
    table.rows[row].cells[cell].text = value


document = Document(TEMPLATE)

for table_index in (0, 21, 24, 27, 29, 34, 37, 41, 43):
    set_cell(document.tables[table_index], 2, 1, "{{UNDER_SECTION}}")

set_cell(document.tables[1], 8, 2, "{{NATURE_OF_ACCIDENT}}")
set_cell(document.tables[38], 3, 3, "{{NATURE_OF_ACCIDENT}}")

form_vi_a = document.tables[31]
death_identity_values = {
    0: "{{DEATH_NAME}} R/o {{DECEASED_ADDRESS}}",
    1: "{{DEATH_FATHER_HUSBAND}}",
    2: "{{DEATH_AGE}}",
    3: "{{DEATH_DATE}}",
    4: "{{DEATH_GENDER}}",
    5: "{{DEATH_MARITAL_STATUS}}",
    6: "{{DEATH_OCCUPATION}}",
    7: "{{DEATH_EMPLOYER}}",
    8: "{{DEATH_INCOME}}",
    9: "{{DEATH_INCOME_TAX}}",
    10: "{{DEATH_SOLE_EARNER}}",
    11: "{{DEATH_MEDICAL_TREATMENT}}",
    12: "{{DEATH_REIMBURSEMENT}}",
}
for row, value in death_identity_values.items():
    set_cell(form_vi_a, row, 18, value)

death_values = {
}
for row, value in death_values.items():
    set_cell(form_vi_a, row, 18, value)

for index, row in enumerate(range(15, 22), start=1):
    set_cell(form_vi_a, row, 2, f"{{{{LEGAL_REP_{index}_NAME}}}}")
    set_cell(form_vi_a, row, 7, f"{{{{LEGAL_REP_{index}_AGE}}}}")
    set_cell(form_vi_a, row, 10, f"{{{{LEGAL_REP_{index}_GENDER}}}}")
    set_cell(form_vi_a, row, 14, f"{{{{LEGAL_REP_{index}_RELATION}}}}")
    set_cell(form_vi_a, row, 18, f"{{{{LEGAL_REP_{index}_MARITAL_STATUS}}}}")

for index, row in enumerate(range(26, 33), start=1):
    set_cell(form_vi_a, row, 0, str(index))
    set_cell(form_vi_a, row, 2, f"{{{{LEGAL_REP_{index}_NAME}}}}")
    set_cell(form_vi_a, row, 3, f"{{{{LEGAL_REP_{index}_CONTACT}}}}")
    set_cell(form_vi_a, row, 13, f"{{{{LEGAL_REP_{index}_ADDRESS}}}}")

for index, row in enumerate(range(35, 41), start=1):
    set_cell(form_vi_a, row, 0, str(index))
    set_cell(form_vi_a, row, 2, f"{{{{CHILD_{index}_NAME}}}}")
    set_cell(form_vi_a, row, 5, f"{{{{CHILD_{index}_SCHOOL_CLASS}}}}")
    set_cell(form_vi_a, row, 12, f"{{{{CHILD_{index}_ANNUAL_FEE}}}}")
    set_cell(form_vi_a, row, 15, f"{{{{CHILD_{index}_EXPENSES}}}}")

injury_values = {
    42: "{{INJURY_NAME}}",
    43: "{{INJURY_FATHER}}",
    44: "{{INJURY_ADDRESS}}",
    45: "{{INJURY_CONTACT}}",
    46: "{{INJURY_AGE}}",
    47: "{{INJURY_GENDER}}",
    48: "{{INJURY_MARITAL_STATUS}}",
    49: "{{INJURY_OCCUPATION}}",
    50: "{{INJURY_EMPLOYER}}",
    51: "{{INJURY_INCOME}}",
    52: "{{INJURY_INCOME_TAX}}",
    53: "{{INJURY_NATURE}}",
    54: "{{INJURY_TREATMENT}}",
    55: "{{INJURY_HOSPITALIZATION}}",
    56: "{{INJURY_SURGERY}}",
    57: "{{INJURY_PERMANENT_DISABILITY}}",
}
for row, value in injury_values.items():
    set_cell(form_vi_a, row, 8, value)

pecuniary_values = {
    74: "{{TREATMENT_EXPENSES}}",
    75: "{{FUTURE_TREATMENT_EXPENSES}}",
    76: "{{CONVEYANCE_EXPENSES}}",
    77: "{{LOSS_OF_INCOME}}",
    78: "{{LOSS_OF_EARNING_CAPACITY}}",
    79: "{{OTHER_PECUNIARY_LOSS}}",
    80: "{{VICTIM_REIMBURSEMENT}}",
    81: "{{PROPERTY_LOSS}}",
    82: "{{ADDITIONAL_INFORMATION}}",
    83: "{{ACCIDENT_BRIEF}}",
    84: "{{COMPENSATION_CLAIMED}}",
}
for row, value in pecuniary_values.items():
    set_cell(form_vi_a, row, 17 if row not in (74, 75, 76) else (4 if row in (74, 75) else 19), value)

form_vii = document.tables[38]
form_vii_values = {
    60: "{{INJURY_NAME}}",
    61: "{{INJURY_AGE}}",
    62: "{{INJURY_OCCUPATION}}",
    63: "{{INJURY_NATURE}}",
    64: "{{INJURY_NATURE}}",
    65: "{{INJURY_NATURE}}",
    66: "{{INJURY_TREATMENT}}",
}
for row, value in form_vii_values.items():
    set_cell(form_vii, row, 3, value)

document.save(TEMPLATE)
